from docx import Document
from typing import Tuple, Dict, Any

class DOCXExtractor:
    async def extract(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        metadata = {'paragraph_count': len(doc.paragraphs)}
        return content, metadata
